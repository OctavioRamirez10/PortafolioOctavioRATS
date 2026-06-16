import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
import tempfile
import shutil

ROOT = Path(__file__).resolve().parents[1]
cv_path = ROOT / 'cv.json'
template_path = ROOT / 'output' / 'plantilla_dark_mode.docx'
photo_path = ROOT / 'portada.png'
out_path = ROOT / 'output' / 'CV_Octavio_Ramirez_ATS_generated.docx'

with open(cv_path, 'r', encoding='utf-8') as f:
    data = json.load(f)

# Load template
if template_path.exists():
    doc = Document(str(template_path))
else:
    doc = Document()

# Insert photo at top if exists
if photo_path.exists():
    # Add a paragraph at top and insert picture
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(str(photo_path), width=Inches(1.5))

# Name and label
name = data['basics'].get('name','')
label = data['basics'].get('label','')

h = doc.add_paragraph()
run = h.add_run(name + '\n')
run.bold = True
run.font.size = Pt(18)
run = h.add_run(label)
run.font.size = Pt(12)

# Contact
contact = doc.add_paragraph()
contact.add_run('Contacto\n').bold = True
basics = data['basics']
contact.add_run(f"Email: {basics.get('email','')}\n")
contact.add_run(f"Teléfono: {basics.get('phone','')}\n")
contact.add_run(f"LinkedIn: {basics.get('profiles',[{}])[0].get('url','')}\n")
contact.add_run(f"GitHub: {basics.get('profiles',[{},{}])[1].get('url','')}\n")
contact.add_run(f"Ubicación: {basics.get('location',{}).get('region','')}\n")

# Profile
doc.add_paragraph('Perfil profesional', style='Heading 2')
summary = basics.get('summary','')
doc.add_paragraph(summary)

# Skills (flatten keywords)
doc.add_paragraph('Habilidades técnicas', style='Heading 2')
skills = []
for s in data.get('skills',[]):
    skills.extend(s.get('keywords',[]))
if skills:
    doc.add_paragraph(', '.join(skills))

# Experience
doc.add_paragraph('Experiencia profesional', style='Heading 2')
for w in data.get('work',[]):
    title = f"{w.get('name','')} — {w.get('position','')}"
    doc.add_paragraph(title, style='Heading 3')
    start = w.get('startDate','')
    end = w.get('endDate','Presente') or 'Presente'
    doc.add_paragraph(f"{start} — {end}")
    if w.get('summary'):
        doc.add_paragraph(w.get('summary'))
    for h in w.get('highlights',[]):
        para = doc.add_paragraph(style='List Bullet')
        para.add_run(h)

# Projects
doc.add_paragraph('Proyectos destacados', style='Heading 2')
for p in data.get('projects',[]):
    doc.add_paragraph(p.get('name',''), style='Heading 3')
    if p.get('description'):
        doc.add_paragraph(p.get('description'))
    for h in p.get('highlights',[]):
        para = doc.add_paragraph(style='List Bullet')
        para.add_run(h)
    if p.get('url'):
        doc.add_paragraph(f"Repo/URL: {p.get('url')}")

# Education
doc.add_paragraph('Educación', style='Heading 2')
for e in data.get('education',[]):
    title = f"{e.get('institution','')} — {e.get('area','')}"
    doc.add_paragraph(title, style='Heading 3')
    start = e.get('startDate','')
    end = e.get('endDate','') or ''
    doc.add_paragraph(f"{start} — {end}")
    for c in e.get('courses',[]):
        para = doc.add_paragraph(style='List Bullet')
        para.add_run(c)

# Certificates
if data.get('certificates'):
    doc.add_paragraph('Certificados', style='Heading 2')
    for c in data.get('certificates'):
        doc.add_paragraph(f"{c.get('name','')} — {c.get('issuer','')} ({c.get('date','')})")

# Languages
if data.get('languages'):
    doc.add_paragraph('Idiomas', style='Heading 2')
    for l in data.get('languages'):
        doc.add_paragraph(f"{l.get('language','')} — {l.get('fluency','')}")

# Save
out_path.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(out_path))
print(f"Saved: {out_path}")
